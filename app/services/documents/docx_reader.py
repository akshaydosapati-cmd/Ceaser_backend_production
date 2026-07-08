from __future__ import annotations

from pathlib import Path

from docx import Document

from app.services.documents.schemas import ExtractedDocument


class DOCXReader:
    def read(self, path: Path) -> ExtractedDocument:
        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return ExtractedDocument(title=path.stem, pages=1, content="\n".join(paragraphs), metadata={"reader": "python-docx"})
